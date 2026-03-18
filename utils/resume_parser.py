import PyPDF2
import docx
import io


SUPPORTED_FORMATS = ["pdf", "docx"]


def extract_text_from_resume(uploaded_file, file_name=None):
    """
    Extracts text from a PDF or Word resume.
    - uploaded_file: Streamlit UploadedFile object or file path string (CLI)
    - file_name: needed when passing file path string in CLI mode
    """
    try:
        # --- Determine file name ---
        if isinstance(uploaded_file, str):
            # CLI mode — uploaded_file is a file path
            name = uploaded_file
        else:
            # Streamlit mode — get name from UploadedFile object
            name = uploaded_file.name

        # --- Get file extension ---
        extension = name.split(".")[-1].lower()

        # --- Check supported format ---
        if extension not in SUPPORTED_FORMATS:
            return None, "⚠️ Unsupported file format. Please upload your resume as a PDF or Word (.docx) file only."

        # --- Route to correct parser ---
        if extension == "pdf":
            text = _parse_pdf(uploaded_file)
        elif extension == "docx":
            text = _parse_docx(uploaded_file)

        # --- Check if text was extracted ---
        if not text or not text.strip():
            return None, "⚠️ Could not extract text from your resume. Make sure it is not a scanned image PDF."

        # --- Validate it looks like a resume ---
        is_valid, error = validate_resume_text(text)
        if not is_valid:
            return None, error

        return text.strip(), None

    except Exception as e:
        return None, f"⚠️ Error reading resume: {str(e)}"


def _parse_pdf(uploaded_file):
    """
    Extracts text from a PDF file.
    """
    if isinstance(uploaded_file, str):
        with open(uploaded_file, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return _extract_pdf_pages(reader)
    else:
        pdf_bytes = io.BytesIO(uploaded_file.read())
        reader = PyPDF2.PdfReader(pdf_bytes)
        return _extract_pdf_pages(reader)


def _extract_pdf_pages(reader):
    """
    Loops through all PDF pages and extracts text.
    """
    full_text = ""
    for page in reader.pages:
        full_text += page.extract_text() + "\n"
    return full_text


def _parse_docx(uploaded_file):
    """
    Extracts text from a Word (.docx) file.
    """
    if isinstance(uploaded_file, str):
        doc = docx.Document(uploaded_file)
    else:
        doc_bytes = io.BytesIO(uploaded_file.read())
        doc = docx.Document(doc_bytes)

    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"
    return full_text


def validate_resume_text(text):
    """
    Basic check to confirm the uploaded file looks like a resume.
    """
    keywords = [
        "experience", "education", "skills",
        "project", "work", "university",
        "college", "internship", "certification"
    ]

    text_lower = text.lower()
    matches = [kw for kw in keywords if kw in text_lower]

    if len(matches) < 2:
        return False, "⚠️ This doesn't look like a resume. Please upload your correct resume file."

    return True, None